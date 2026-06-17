import os
import json
import io
import requests
from datetime import datetime
from urllib.parse import urljoin

from bs4 import BeautifulSoup
from fastapi import FastAPI
import gspread
from google.oauth2.service_account import Credentials

import PyPDF2
from docx import Document
import openpyxl


app = FastAPI(title="AI Tender Agent Cargo V22 Document Checklist")

BOT_TOKEN = os.getenv("BOT_TOKEN")
CHAT_ID = os.getenv("CHAT_ID")
GOOGLE_SHEET_ID = os.getenv("GOOGLE_SHEET_ID")


SERVICE_PHRASES = [
    "перевозка грузов", "перевозке грузов", "перевозки грузов",
    "грузоперевоз", "грузовые перевозки",
    "доставка грузов", "доставка груза", "доставка товара автотранспортом",
    "транспортные услуги", "оказание транспортных услуг",
    "услуги транспорта", "услуги по перевозке", "услуги перевозки",
    "транспортно-экспедиционные услуги", "транспортно экспедиционные услуги",
    "экспедиторские услуги", "логистические услуги",
    "международные автомобильные перевозки", "международная перевозка",
    "международные перевозки", "автомобильные перевозки грузов",
    "cargo transportation", "freight transportation", "logistics service",
    "transport service", "delivery service",
    "yuk tashish", "yuklarni tashish", "yuk tashish xizmati",
    "yuklarni tashish xizmati", "yuk tashuvchi xizmatlari",
    "yuklarni tashuvchi", "transport xizmati", "transport xizmatlari",
    "logistika xizmati", "logistika xizmatlari",
    "ekspeditorlik xizmati", "ekspeditorlik xizmatlari",
    "xalqaro yuk tashish", "xalqaro tashuv", "xalqaro tashish",
    "avtotransport xizmati", "avtotransport xizmatlari",
]

SUPPORT_WORDS = [
    "перевоз", "груз", "достав", "логист", "экспед",
    "услуг", "хизмат", "xizmat", "xizmatlari", "xizmati",
    "cargo", "freight", "delivery", "logistics",
    "yuk", "yuklarni", "tashish", "tashuvchi", "tashuv",
    "xalqaro", "transport", "avtotransport",
]

HARD_BAD_WORDS = [
    "закупка", "поставка", "приобретение", "купить", "сотиб олиш",
    "машин", "машина", "автомобиль грузовой", "грузовой автомобиль",
    "шасси", "ось", "двигател", "мотор", "запчаст", "запасн",
    "инструмент", "набор инструментов", "оборудован", "техника",
    "смесительно-заряд", "зарядных машин", "спецтехника",
    "автокран", "погрузчик", "экскаватор", "трактор",
    "арматур", "бетон", "цемент", "лаборатор", "мебел", "пленк",
    "стретч", "консультац", "технадзор", "строительств", "ремонт",
    "канцеляр", "компьютер", "принтер", "медицин", "питание",
    "продукт", "одежд", "обув", "электро", "юридическ", "аудит",
    "страхован", "охрана", "дезинфек", "дезинсек", "deratiz",
    "овқат", "еда", "питания", "payvandlash",
    "метал конструкц", "metal konstruksiya",
    "yo‘li", "yo'li", "yoʻli", "avtomobil yo",
    "yer uchastkasi", "master-reja", "baholash",
    "service area", "service point", "проект", "лойиҳа",
    "автоматизац", "кредит", "кредитного конвейера", "банк", "банковск",
    "финанс", "программ", "программного обеспечения", "по для",
    "система принятия решений", "принятия решений", "внедрение системы",
    "внедрению системы", "it", "информационн", "цифров",
    "software", "sap", "oracle", "crm", "erp", "сервер",
    "лиценз", "лицензи", "разработк", "поддержк программ",
    "интеграц", "модернизац системы", "сопровождение системы",
    "кибербезопас", "телеком", "интернет", "сайт", "портал",
]

SOFT_BAD_WORDS = [
    "товар", "материал", "изделие", "деталь", "агрегат", "насос",
    "кабель", "труба", "краска", "масло"
]

QUALITY_BAD_CONTEXT_WORDS = [
    "инерт материал", "инертные материалы", "щебень", "песок", "гравий",
    "бетон", "цемент", "арматура", "кирпич", "строительный материал",
    "қум", "шағал", "inert material", "qum", "shag'al",
    "утилизация", "чиқинди", "отход", "maishiy chiqindi",
    "qattiq chiqindi", "мусор", "вывоз мусора",
    "типограф", "печать", "газет", "bosmaxona", "chop etish",
    "овқат", "питание", "еда", "медицин", "лаборатор",
    "консультац", "технадзор", "аудит", "юридическ",
    "закупка", "поставка", "приобретение", "сотиб олиш",
    "запчаст", "запасн", "ось", "шасси", "двигател", "оборудован",
    "смесительно-заряд", "инструмент", "набор инструментов",
    "автоматизац", "кредит", "банк", "программ", "система принятия решений",
]

QUALITY_GOOD_CONTEXT_WORDS = [
    "перевозка грузов", "перевозке грузов", "услуга перевозка грузов",
    "транспортно-экспедитор", "экспедиторские услуги",
    "логистические услуги", "международная перевозка", "международные перевозки",
    "мультимодальная", "мультимодальные", "негабарит", "тяжеловес",
    "yuk tashish", "yuklarni tashish", "yuk tashish xizmati",
    "xalqaro yuk tashish", "xalqaro", "logistika xizmatlari",
    "ekspeditorlik", "multimodal", "gabarit bo", "og'ir vazn", "og’ir vazn",
]

QUALITY_STRONG_GOOD_WORDS = [
    "xalqaro", "международ", "multimodal", "мультимод",
    "негабарит", "тяжеловес", "gabarit bo", "og'ir vazn", "og’ir vazn",
    "экспедитор", "logistika",
]

BAD_URL_PARTS = [
    "register", "login", "logout", "signin", "signup", "cabinet", "profile",
    "account", "user", "my", "add.html", "/add", "create", "invited",
    "invitation", "english", "/en/", "news", "blog", "faq", "help",
    "contact", "about", "rules", "terms", "privacy", "advertising",
    "banner", "calendar", "archive", "feedback", "javascript:", "mailto:", "tel:",
]

BAD_TITLE_WORDS = [
    "регистрация", "зарегистрироваться", "войти", "выход",
    "стать заказчиком", "стать поставщиком", "english", "русский", "ўзбекча",
    "приглашение", "мои заявки", "моих заявок", "вопрос", "ответ",
    "помощь", "контакты", "написать нам письмо", "о сайте", "правила",
    "дата публикации", "личный кабинет", "кабинет", "профиль",
]


def clean_text(text):
    return " ".join((text or "").replace("\n", " ").replace("\t", " ").split())


def normalize_text(text):
    text = clean_text(text).lower()
    text = text.replace("’", "'").replace("‘", "'").replace("ʻ", "'").replace("`", "'")
    text = text.replace("ё", "е")
    return text


def get_headers(json_mode=False):
    return {
        "User-Agent": "Mozilla/5.0",
        "Accept": "application/json,text/plain,*/*" if json_mode else "text/html,application/xhtml+xml,*/*",
        "Accept-Language": "ru-RU,ru;q=0.9,uz;q=0.8,en;q=0.7",
        "Content-Type": "application/json" if json_mode else "text/html",
    }


def looks_like_bad_url(url):
    url = (url or "").lower()
    return any(part in url for part in BAD_URL_PARTS)


def filter_reason(title, url):
    t = normalize_text(title)

    if not title or not url:
        return "empty_title_or_url"
    if looks_like_bad_url(url):
        return "bad_url"
    if title.isdigit():
        return "only_digits"
    if len(title.split()) < 2:
        return "too_short"
    if len(t) < 10:
        return "too_short_text"

    for bad in BAD_TITLE_WORDS:
        if bad in t:
            return "bad_title_word:" + bad

    for bad in HARD_BAD_WORDS:
        if bad in t:
            return "hard_bad_word:" + bad

    for phrase in SERVICE_PHRASES:
        if phrase in t:
            return "accepted_service_phrase:" + phrase

    support_count = sum(1 for word in SUPPORT_WORDS if word in t)
    has_cargo = any(x in t for x in ["груз", "yuk", "cargo", "freight"])
    has_service = any(x in t for x in ["услуг", "xizmat", "service"])
    has_action = any(x in t for x in ["перевоз", "достав", "tashish", "tashuv", "delivery", "transportation"])

    if support_count >= 3 and has_cargo and has_service and has_action:
        if any(x in t for x in SOFT_BAD_WORDS):
            return "soft_bad_but_possible"
        return "accepted_support_combo"

    return "not_service_cargo"


def is_real_cargo_tender(title, url):
    return filter_reason(title, url).startswith("accepted_")


def get_sheet():
    raw_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    if not raw_json:
        raise Exception("GOOGLE_SERVICE_ACCOUNT_JSON is empty")

    info = json.loads(raw_json)
    creds = Credentials.from_service_account_info(
        info,
        scopes=[
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive",
        ],
    )
    client = gspread.authorize(creds)
    return client.open_by_key(GOOGLE_SHEET_ID).worksheet("Тендеры")


def send_telegram(text):
    if not BOT_TOKEN or not CHAT_ID:
        return False

    try:
        response = requests.post(
            f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage",
            json={"chat_id": CHAT_ID, "text": text[:3900]},
            timeout=8,
        )
        return response.status_code == 200
    except Exception as e:
        print("TELEGRAM ERROR:", e)
        return False


def make_key(site, title, url):
    return f"{site}:{normalize_text(title)[:120]}:{url}".lower()


def add_tender(tenders, seen, site, title, url):
    title = clean_text(title)
    url = clean_text(url)

    if not title or not url:
        return

    key = make_key(site, title, url)
    if key in seen:
        return

    if not is_real_cargo_tender(title, url):
        return

    seen.add(key)
    tenders.append({
        "site": site,
        "title": title[:300],
        "url": url,
        "reason": filter_reason(title, url),
    })


def add_raw_candidate(candidates, seen, site, title, url):
    title = clean_text(title)
    url = clean_text(url)

    if not title or not url:
        return

    key = make_key(site, title, url)
    if key in seen:
        return

    seen.add(key)
    candidates.append({
        "site": site,
        "title": title[:300],
        "url": url,
        "accepted": is_real_cargo_tender(title, url),
        "reason": filter_reason(title, url),
    })


def collect_links(base_url, pages, site, limit=10, raw=False):
    tenders = []
    candidates = []
    seen = set()

    for page_url in pages[:limit]:
        try:
            r = requests.get(page_url, headers=get_headers(), timeout=12)
            print(f"{site} PAGE:", page_url)
            print(f"{site} STATUS:", r.status_code)
            print(f"{site} SIZE:", len(r.text))

            if r.status_code != 200:
                continue

            soup = BeautifulSoup(r.text, "html.parser")

            for a in soup.find_all("a"):
                title = clean_text(a.get_text(" ", strip=True))
                href = a.get("href")
                if not title or not href:
                    continue

                url = urljoin(base_url, href)

                if raw:
                    add_raw_candidate(candidates, seen, site, title, url)
                else:
                    add_tender(tenders, seen, site, title, url)

            for tag in soup.find_all(["div", "tr", "li", "article", "section"]):
                text = clean_text(tag.get_text(" ", strip=True))
                if 20 <= len(text) <= 300:
                    if raw:
                        add_raw_candidate(candidates, seen, site, text, page_url)
                    else:
                        add_tender(tenders, seen, site, text, page_url)

        except Exception as e:
            print(f"{site} HTML ERROR:", e)

    return candidates if raw else tenders


def flatten_items(data):
    items = []

    if isinstance(data, list):
        for x in data:
            if isinstance(x, dict):
                items.append(x)
                items.extend(flatten_items(x))
            elif isinstance(x, list):
                items.extend(flatten_items(x))

    elif isinstance(data, dict):
        for v in data.values():
            if isinstance(v, dict):
                items.append(v)
                items.extend(flatten_items(v))
            elif isinstance(v, list):
                items.extend(flatten_items(v))

    return items


def item_url(item, base_url, site):
    for key in ["url", "link", "href"]:
        if item.get(key):
            return urljoin(base_url, str(item.get(key)))

    lot_id = (
        item.get("id")
        or item.get("lotId")
        or item.get("number")
        or item.get("lotNumber")
        or item.get("procedureId")
        or item.get("display_no")
    )

    if site == "UZEX" and lot_id:
        return f"https://etender.uzex.uz/lot/{lot_id}"

    if site == "XT-Xarid" and lot_id:
        return f"https://xt-xarid.uz/tender/{lot_id}"

    return base_url


def item_title(item):
    keys = [
        "title", "name", "lotName", "productName", "subject",
        "description", "descriptionRu", "nameRu", "nameUz",
        "goodsName", "serviceName", "procedureName",
        "category_name", "company_name",
    ]

    parts = []

    for key in keys:
        value = item.get(key)
        if value and str(value).lower() not in ["tender", "none", "null"]:
            parts.append(str(value))

    meta = item.get("meta")
    if isinstance(meta, dict):
        for gm in meta.get("good_maps", []):
            if isinstance(gm, dict):
                if gm.get("name"):
                    parts.append(str(gm.get("name")))

                category = gm.get("category")
                if isinstance(category, dict) and category.get("title"):
                    parts.append(str(category.get("title")))

        if meta.get("company_name"):
            parts.append(str(meta.get("company_name")))

    texts = [str(v) for v in item.values() if isinstance(v, str) and len(v) > 8]
    parts.extend(texts[:3])

    return clean_text(" | ".join(parts))


def parse_tenderweek(raw=False):
    base_url = "https://www.tenderweek.com/"
    pages = [base_url]

    search_queries = [
        "перевозка грузов",
        "транспортные услуги",
        "логистические услуги",
        "транспортно-экспедиционные услуги",
        "доставка грузов",
    ]

    for page in range(2, 10):
        pages.append(f"{base_url}?page={page}")

    for q in search_queries:
        pages.append(f"{base_url}?search={requests.utils.quote(q)}")
        pages.append(f"{base_url}?q={requests.utils.quote(q)}")

    return collect_links(base_url, pages, "Tenderweek", limit=20, raw=raw)


def parse_uzex_api(raw=False):
    url = "https://apietender.uzex.uz/api/common/TradeList"
    base_url = "https://etender.uzex.uz/"
    tenders = []
    candidates = []
    seen = set()

    payloads = [
        {"TypeId": 1, "From": 1, "To": 200, "System_Id": 0},
        {"TypeId": 2, "From": 1, "To": 200, "System_Id": 0},
        {"TypeId": 3, "From": 1, "To": 200, "System_Id": 0},
    ]

    for payload in payloads:
        try:
            r = requests.post(url, headers=get_headers(json_mode=True), json=payload, timeout=15)

            print("UZEX API STATUS:", r.status_code)
            print("UZEX API SIZE:", len(r.text))
            print("UZEX API TYPE:", r.headers.get("content-type", ""))

            if r.status_code != 200:
                continue

            data = r.json()
            items = flatten_items(data)
            print("UZEX API ITEMS:", len(items))

            for item in items:
                if not isinstance(item, dict):
                    continue

                title = item_title(item)
                tender_url = item_url(item, base_url, "UZEX")

                if raw:
                    add_raw_candidate(candidates, seen, "UZEX", title, tender_url)
                else:
                    add_tender(tenders, seen, "UZEX", title, tender_url)

        except Exception as e:
            print("UZEX API ERROR:", e)

    return candidates if raw else tenders


def parse_xt_xarid_api(raw=False):
    url = "https://api.xt-xarid.uz/rpc"
    base_url = "https://xt-xarid.uz/"
    tenders = []
    candidates = []
    seen = set()

    offsets = [0, 100, 200, 300]

    for offset in offsets:
        payload = {
            "id": 1,
            "jsonrpc": "2.0",
            "method": "ref",
            "params": {
                "ref": "ref_tender_public",
                "op": "read",
                "limit": 100,
                "offset": offset,
                "filters": {},
            },
        }

        try:
            r = requests.post(url, headers=get_headers(json_mode=True), json=payload, timeout=15)

            print("XT-Xarid API STATUS:", r.status_code)
            print("XT-Xarid API SIZE:", len(r.text))
            print("XT-Xarid API TYPE:", r.headers.get("content-type", ""))

            if r.status_code != 200:
                continue

            data = r.json()
            items = flatten_items(data)
            print("XT-Xarid API ITEMS:", len(items))

            for item in items:
                if not isinstance(item, dict):
                    continue

                title = item_title(item)
                tender_url = item_url(item, base_url, "XT-Xarid")

                if raw:
                    add_raw_candidate(candidates, seen, "XT-Xarid", title, tender_url)
                else:
                    add_tender(tenders, seen, "XT-Xarid", title, tender_url)

        except Exception as e:
            print("XT-Xarid API ERROR:", e)

    return candidates if raw else tenders


def parse_xt_xarid(raw=False):
    return parse_xt_xarid_api(raw=raw)


def parse_uzex(raw=False):
    return parse_uzex_api(raw=raw)


def extract_lot_id_from_url(url):
    if not url:
        return None

    url = str(url).strip().rstrip("/")
    parts = url.split("/")

    for part in reversed(parts):
        if part.isdigit():
            return part

    return None


def safe_json_loads(raw, default=None):
    if default is None:
        default = []

    if not raw:
        return default

    try:
        return json.loads(raw)
    except Exception:
        return default


def get_uzex_lot_data(lot_id):
    api_url = f"https://apietender.uzex.uz/api/common/GetTrade/{lot_id}/0"
    r = requests.get(api_url, headers=get_headers(json_mode=True), timeout=20)
    r.raise_for_status()
    return r.json()


def get_uzex_trade(lot_id):
    return get_uzex_lot_data(lot_id)


def extract_budget_products(trade):
    raw = trade.get("budget_products")
    if not raw:
        return []

    try:
        return json.loads(raw)
    except Exception:
        return []


def short_requirements_from_trade(trade):
    parts = []

    budget_products = safe_json_loads(trade.get("budget_products"), [])
    if budget_products:
        p = budget_products[0]
        desc = clean_text(p.get("Description", ""))
        delivery = p.get("Delivery_Term")
        category = clean_text(p.get("Category_Name", ""))
        product_name = clean_text(p.get("Product_Name", ""))

        if product_name:
            parts.append("Предмет: " + product_name)
        if desc:
            parts.append(desc[:350])
        if category:
            parts.append("Категория: " + category)
        if delivery:
            parts.append(f"Срок оказания: {delivery} дней")

    q_fields = trade.get("js_qualification_fields") or []
    if isinstance(q_fields, list) and q_fields:
        q_short = []
        for q in q_fields[:8]:
            name = clean_text(q.get("name", "") or q.get("title", "") or q.get("label", ""))
            if name:
                q_short.append(name[:120])
        if q_short:
            parts.append("Квалификация: " + "; ".join(q_short))

    return clean_text(" | ".join(parts))[:1200]


def quality_decision_from_trade(trade, title=""):
    budget_products = safe_json_loads(trade.get("budget_products"), [])
    description = ""
    product_name = ""
    category_name = ""

    if budget_products:
        product_name = clean_text(budget_products[0].get("Product_Name", ""))
        description = clean_text(budget_products[0].get("Description", ""))
        category_name = clean_text(budget_products[0].get("Category_Name", ""))

    text = normalize_text(" ".join([
        title or "",
        product_name,
        description,
        category_name,
        trade.get("customer_name") or "",
        trade.get("technical_description") or "",
    ]))

    bad_hits = [w for w in QUALITY_BAD_CONTEXT_WORDS if w in text]
    good_hits = [w for w in QUALITY_GOOD_CONTEXT_WORDS if w in text]
    strong_hits = [w for w in QUALITY_STRONG_GOOD_WORDS if w in text]

    if strong_hits and any(w in text for w in ["yuk tashish", "перевоз", "transport", "логист", "экспедитор"]):
        return {
            "logistics": "Да",
            "priority": "Очень высокий",
            "risk": "Средний",
            "win_chance": "Высокий",
            "reason": "Сильный профиль логистики: " + ", ".join(strong_hits[:4]),
            "decision": "Участвовать: высокий приоритет, международная/сложная логистика, нужны партнёры и расчёт себестоимости",
        }

    if bad_hits and not strong_hits:
        return {
            "logistics": "Нет",
            "priority": "Низкий",
            "risk": "Высокий",
            "win_chance": "Низкий",
            "reason": "Непрофильный контекст: " + ", ".join(bad_hits[:5]),
            "decision": "Отказаться: непрофильная закупка или не транспортная услуга для Trans Ocean Logistics",
        }

    if good_hits:
        return {
            "logistics": "Да",
            "priority": "Высокий",
            "risk": "Средний",
            "win_chance": "Средний",
            "reason": "Профильная логистическая услуга: " + ", ".join(good_hits[:5]),
            "decision": "Рассмотреть участие: профильная логистическая услуга",
        }

    if any(w in text for w in ["transport", "транспорт", "avtotransport", "yuk"]):
        return {
            "logistics": "Сомнительно",
            "priority": "Средний",
            "risk": "Средний",
            "win_chance": "Средний",
            "reason": "Есть транспортные слова, но требуется ручная проверка",
            "decision": "Изучить: требуется ручная проверка условий",
        }

    return {
        "logistics": "Нет",
        "priority": "Низкий",
        "risk": "Высокий",
        "win_chance": "Низкий",
        "reason": "Нет достаточных признаков логистической услуги",
        "decision": "Отказаться: недостаточно признаков профильной логистики",
    }


def detect_required_documents_from_text(title, text, trade=None):
    """
    Cargo V22:
    Автоматически определяет документы, которые нужно подготовить сотрудникам.
    Работает по API-полям UZEX + тексту ТЗ/PDF/DOCX, если документы удалось прочитать.
    """
    t = normalize_text((title or "") + " " + (text or ""))

    checklist = []
    warnings = []
    responsible = []

    def add_doc(name, reason=""):
        if name not in [x["document"] for x in checklist]:
            checklist.append({
                "document": name,
                "reason": reason,
                "status": "Подготовить"
            })

    def add_warning(w):
        if w not in warnings:
            warnings.append(w)

    add_doc("Коммерческое предложение / ценовое предложение", "Основной документ для участия")
    add_doc("Реквизиты компании", "Нужно для договора и заявки")
    add_doc("Свидетельство о регистрации компании", "Стандартный регистрационный документ")
    add_doc("Устав компании", "Часто требуется для подтверждения полномочий")
    add_doc("Доверенность на подписанта или приказ директора", "Если заявку подписывает не директор")
    add_doc("Паспортные данные / ID подписанта", "Для подтверждения полномочий")
    add_doc("Информация об опыте аналогичных перевозок", "Для подтверждения квалификации")
    add_doc("Список транспорта / данные по машинам", "Для подтверждения возможности выполнить перевозку")
    add_doc("Данные по водителям", "Если требуется допуск водителей к перевозке")
    add_doc("Контакты ответственного менеджера", "Для связи с заказчиком")

    if any(w in t for w in ["лиценз", "license", "ruxsatnoma", "рухсатнома"]):
        add_doc("Лицензия / разрешение на транспортную деятельность", "В ТЗ есть признак требования лицензии")
        add_warning("Проверить наличие лицензии или разрешения")

    if any(w in t for w in ["сертификат", "certificate", "sertifikat"]):
        add_doc("Сертификаты / разрешительные документы", "В документах есть упоминание сертификатов")

    if any(w in t for w in ["налог", "задолж", "qarzdorlik", "soliq"]):
        add_doc("Справка об отсутствии налоговой задолженности", "Есть признак требования по налоговой чистоте")
        add_warning("Проверить налоговую задолженность до подачи")

    if any(w in t for w in ["банк гаран", "банковская гарантия", "bank kafolat", "kafolat"]):
        add_doc("Банковская гарантия", "В документах есть признак банковской гарантии")
        add_warning("Срочно проверить сумму и срок банковской гарантии")

    if any(w in t for w in ["обеспечение заявки", "залог", "deposit", "zakalat", "garov"]):
        add_doc("Обеспечение заявки / залог", "Есть признак обеспечения заявки")
        add_warning("Проверить размер обеспечения заявки")

    if any(w in t for w in ["страхов", "insurance", "sug'urta", "sugurta"]):
        add_doc("Страховой полис груза / CMR страхование", "Есть признак требования страхования")
        add_warning("Проверить страховое покрытие груза")

    if any(w in t for w in ["cmr", "смр"]):
        add_doc("CMR накладная / международные транспортные документы", "Есть признак международной перевозки")
        add_doc("Документы на транспорт для международной перевозки", "CMR требует корректных данных транспорта")

    if any(w in t for w in ["тамож", "customs", "bojxona"]):
        add_doc("Таможенные документы / данные для оформления", "Есть признак таможенного оформления")
        add_warning("Проверить, кто отвечает за таможенное оформление")

    if any(w in t for w in ["международ", "xalqaro", "international"]):
        add_doc("Документы для международной перевозки", "Тендер похож на международную перевозку")
        add_doc("Договоры/акты по прошлым международным перевозкам", "Для подтверждения опыта")
        add_warning("Проверить маршрут, погранпереходы, разрешения и транзитные страны")

    if any(w in t for w in ["негабарит", "тяжеловес", "og'ir vazn", "og’ir vazn", "gabarit bo"]):
        add_doc("Разрешения на негабаритный / тяжеловесный груз", "Есть признак сложной перевозки")
        add_doc("Схема крепления груза / план перевозки", "Для сложной логистики")
        add_doc("Фото/техпаспорт спецтранспорта", "Для подтверждения технической возможности")
        add_warning("Высокий риск: негабарит/тяжеловес требует спецразрешений и расчёта маршрута")

    if any(w in t for w in ["рефриж", "холод", "температур", "refrigerator"]):
        add_doc("Документы на рефрижераторный транспорт", "Есть температурный режим")
        add_doc("Подтверждение температурного режима", "Может потребоваться для груза")
        add_warning("Проверить температурный режим и ответственность за порчу груза")

    if any(w in t for w in ["опасн", "adr", "hazard", "xavfli"]):
        add_doc("ADR разрешение / документы для опасного груза", "Есть признак опасного груза")
        add_doc("Допуск водителя ADR", "Для опасного груза")
        add_warning("Опасный груз: требуется отдельная проверка допуска и страховки")

    if any(w in t for w in ["договор", "контракт", "contract", "shartnoma"]):
        add_doc("Подписанный проект договора / протокол разногласий", "Есть проект договора")

    if any(w in t for w in ["техническое задание", "тз", "technical", "texnik topshiriq"]):
        add_doc("Подтверждение соответствия техническому заданию", "Нужно ответить по требованиям ТЗ")

    if any(w in t for w in ["срок", "delivery", "muddat", "календар"]):
        add_doc("График оказания услуг / план перевозки", "Есть требования по срокам")

    if any(w in t for w in ["тендерная комиссия", "квалификац", "qualification", "malaka"]):
        add_doc("Квалификационная анкета участника", "Есть квалификационные требования")

    if trade:
        q_fields = trade.get("js_qualification_fields") or []
        if isinstance(q_fields, list) and q_fields:
            add_doc("Ответы на квалификационные вопросы UZEX", "В API есть квалификационные поля")
            for q in q_fields[:10]:
                if isinstance(q, dict):
                    name = clean_text(q.get("name", "") or q.get("title", "") or q.get("label", ""))
                    if name:
                        add_doc("Квалификационный документ: " + name[:80], "Из поля квалификации UZEX")

    responsible = [
        {"role": "Тендерный менеджер", "task": "Собрать документы и проверить дедлайн"},
        {"role": "Логист", "task": "Проверить маршрут, транспорт, тоннаж, сроки"},
        {"role": "Бухгалтерия", "task": "Проверить налоги, оплату, обеспечение заявки"},
        {"role": "Директор", "task": "Принять решение участвовать / не участвовать"},
    ]

    return {
        "required_documents": checklist,
        "warnings": warnings,
        "responsible_tasks": responsible,
        "documents_count": len(checklist),
        "summary": "; ".join([x["document"] for x in checklist[:12]])
    }


def extract_route_and_transport(title, text):
    t = normalize_text(title + " " + text)

    route_hints = []
    transport_hints = []

    route_words = [
        "ташкент", "toshkent", "узбекистан", "o'zbekiston", "ўзбекистон",
        "казахстан", "qozog", "китай", "xitoy", "россия", "moskva", "москва",
        "хоргос", "xorgos", "яллама", "yallama", "тяньцзинь", "tyantszin",
        "самарканд", "samarqand", "бухара", "buxoro", "андижан", "andijon",
    ]

    for w in route_words:
        if w in t and w not in route_hints:
            route_hints.append(w)

    if "рефриж" in t:
        transport_hints.append("Рефрижератор")
    if "тент" in t:
        transport_hints.append("Тент")
    if "изотерм" in t:
        transport_hints.append("Изотерм")
    if "контейнер" in t:
        transport_hints.append("Контейнер")
    if "низкорам" in t or "негабарит" in t or "тяжеловес" in t:
        transport_hints.append("Низкорамный трал / спецтранспорт")
    if "фура" in t or "20 тонн" in t or "22 тонн" in t or "еврофура" in t:
        transport_hints.append("Фура 20-22 тонн")

    return {
        "route_hints": route_hints[:10],
        "transport_hints": transport_hints[:10],
        "route_text": ", ".join(route_hints[:10]) if route_hints else "Маршрут не найден в API/ТЗ",
        "transport_text": ", ".join(transport_hints[:10]) if transport_hints else "Тип транспорта не найден в API/ТЗ",
    }


def smart_api_analysis_from_trade(trade, title=""):
    budget_products = extract_budget_products(trade)
    product_name = ""
    description = ""
    category_name = ""
    delivery_term = ""

    if budget_products:
        first = budget_products[0]
        product_name = clean_text(first.get("Product_Name", ""))
        description = clean_text(first.get("Description", ""))
        category_name = clean_text(first.get("Category_Name", ""))
        delivery_term = first.get("Delivery_Term", "") or ""

    q_fields = trade.get("js_qualification_fields") or []
    q_texts = []
    if isinstance(q_fields, list):
        for q in q_fields:
            if isinstance(q, dict):
                for key in ["name", "title", "description", "label"]:
                    value = q.get(key)
                    if value:
                        q_texts.append(str(value))

    full_text_raw = " ".join([
        title or "",
        str(trade.get("customer_name") or ""),
        product_name,
        description,
        category_name,
        str(trade.get("technical_description") or ""),
        " ".join(q_texts),
    ])

    text = normalize_text(full_text_raw)

    start_cost = trade.get("start_cost") or 0
    try:
        amount = float(start_cost)
    except Exception:
        amount = 0.0

    payment_type = clean_text(trade.get("payment_type_name", ""))
    term_payment_days = trade.get("term_payment_days", "") or ""

    score = 40
    reasons = []
    risks = []

    def add(points, reason):
        nonlocal score
        score += points
        reasons.append(reason)

    def sub(points, reason):
        nonlocal score
        score -= points
        risks.append(reason)

    if any(w in text for w in ["xalqaro", "международ"]):
        add(25, "международная перевозка / xalqaro yo‘nalish")
    if any(w in text for w in ["logistika", "логист"]):
        add(18, "логистическая услуга")
    if any(w in text for w in ["экспедитор", "ekspeditor"]):
        add(18, "экспедиторская услуга")
    if any(w in text for w in ["перевозка грузов", "перевозке грузов", "yuk tashish", "yuklarni tashish"]):
        add(20, "перевозка грузов")
    if any(w in text for w in ["мультимод", "multimodal", "негабарит", "тяжеловес", "og'ir vazn", "og’ir vazn", "gabarit bo"]):
        add(18, "сложная/дорогая логистика")
    if any(w in normalize_text(payment_type) for w in ["олдиндан", "предоплат", "аванс"]):
        add(10, "есть предоплата")
    if amount >= 1_000_000_000:
        add(20, "крупная сумма больше 1 млрд UZS")
    elif amount >= 500_000_000:
        add(14, "сумма больше 500 млн UZS")
    elif amount >= 100_000_000:
        add(8, "сумма от 100 млн UZS")
    elif amount and amount < 50_000_000:
        sub(10, "маленькая сумма для международной логистики")

    if str(delivery_term).isdigit() and int(delivery_term) >= 180:
        add(8, "длинный срок договора")

    if any(w in text for w in ["погруз", "разгруз", "yuklash", "tushirish", "склад", "ombor"]):
        sub(8, "есть признаки погрузочно-складских работ, не чистая перевозка")
    if any(w in text for w in QUALITY_BAD_CONTEXT_WORDS):
        sub(25, "есть непрофильные слова из стоп-листа")
    if any(w in text for w in ["банк", "кредит", "программ", "автоматизац", "оборудован", "ремонт", "строительств"]):
        sub(30, "похоже на непрофильную закупку")

    score = max(0, min(100, score))

    if score >= 80:
        priority = "Очень высокий"
        win_chance = "Высокий"
        decision = "Участвовать: высокий приоритет. Подготовить расчёт себестоимости и проверить ТЗ вручную на UZEX."
    elif score >= 65:
        priority = "Высокий"
        win_chance = "Средний/Высокий"
        decision = "Рассмотреть участие: профильный тендер, нужна проверка маршрута, транспорта и документов."
    elif score >= 45:
        priority = "Средний"
        win_chance = "Средний"
        decision = "Изучить вручную: есть признаки логистики, но нужна проверка деталей."
    else:
        priority = "Низкий"
        win_chance = "Низкий"
        decision = "Скорее отказаться: недостаточно профильных признаков или высокий риск мусорного лота."

    risk = "Средний"
    if risks and score < 65:
        risk = "Высокий"
    elif risks:
        risk = "Средний"
    elif score >= 80:
        risk = "Низкий/Средний"

    logistics = "Да" if score >= 55 else "Сомнительно" if score >= 40 else "Нет"

    doc_note = "V22: агент формирует чек-лист документов по API/ТЗ. Если PDF/DOCX закрыт, открыть документы вручную на странице лота."

    requirements_short = short_requirements_from_trade(trade)
    reason_text = "; ".join(reasons[:6]) if reasons else "нет сильных положительных факторов"
    risk_text = "; ".join(risks[:5]) if risks else "критических рисков по API не найдено"

    docs = detect_required_documents_from_text(title, full_text_raw, trade)
    route_transport = extract_route_and_transport(title, full_text_raw)

    return {
        "score": score,
        "priority": priority,
        "win_chance": win_chance,
        "risk": risk,
        "logistics": logistics,
        "decision": decision,
        "reasons": reasons,
        "risks": risks,
        "reason_text": reason_text,
        "risk_text": risk_text,
        "document_note": doc_note,
        "requirements_short": requirements_short,
        "api_text_used": clean_text(" | ".join([product_name, description, category_name, " ".join(q_texts[:5])]))[:1200],
        "required_documents": docs["required_documents"],
        "document_warnings": docs["warnings"],
        "responsible_tasks": docs["responsible_tasks"],
        "documents_summary": docs["summary"],
        "documents_count": docs["documents_count"],
        "route_text": route_transport["route_text"],
        "transport_text": route_transport["transport_text"],
    }


def smart_api_analysis_for_url(site, title, url):
    if site != "UZEX":
        return {
            "score": 50,
            "priority": "Средний",
            "win_chance": "Требуется проверка",
            "risk": "Средний",
            "logistics": "Сомнительно",
            "decision": "Проверить вручную: для этого источника нет глубокого UZEX API-анализа.",
            "reason_text": "источник не UZEX",
            "risk_text": "нет API-деталей",
            "document_note": "Документы нужно открыть вручную.",
            "requirements_short": "",
            "required_documents": detect_required_documents_from_text(title, "", None)["required_documents"],
            "document_warnings": [],
            "responsible_tasks": [],
            "documents_summary": "",
            "documents_count": 0,
            "route_text": "Маршрут не найден",
            "transport_text": "Тип транспорта не найден",
        }

    lot_id = extract_lot_id_from_url(url)
    if not lot_id:
        return {
            "score": 0,
            "priority": "Низкий",
            "win_chance": "Низкий",
            "risk": "Высокий",
            "logistics": "Нет",
            "decision": "Не удалось определить ID лота.",
            "reason_text": "нет lot_id",
            "risk_text": "невозможно получить API-данные",
            "document_note": "",
            "requirements_short": "",
            "required_documents": [],
            "document_warnings": ["Не удалось определить ID лота"],
            "responsible_tasks": [],
            "documents_summary": "",
            "documents_count": 0,
            "route_text": "Маршрут не найден",
            "transport_text": "Тип транспорта не найден",
        }

    trade = get_uzex_lot_data(lot_id)
    return smart_api_analysis_from_trade(trade, title)


EXTRA_SHEET_COLUMNS = [
    "Заказчик",
    "Сумма",
    "Валюта",
    "Оплата",
    "Срок оплаты",
    "Срок оказания услуг",
    "Маршрут",
    "Тип транспорта",
    "Требования",
    "Документы нужны",
    "Предупреждения по документам",
    "Ответственные задачи",
    "Рекомендация AI",
]

TENDER_MANAGER_COLUMNS = [
    "Подавать",
    "Ответственный",
    "Дата решения",
    "Статус участия",
    "Комментарий директора",
]


def ensure_sheet_columns():
    sheet = get_sheet()
    headers = sheet.row_values(1)

    if not headers:
        headers = [
            "Дата",
            "Отправил",
            "Ссылка",
            "Источник",
            "Статус",
            "Приоритет",
            "AI анализ",
            "Комментарий",
        ]
        sheet.update("A1:H1", [headers])

    current_headers = sheet.row_values(1)
    added = []

    for col_name in EXTRA_SHEET_COLUMNS:
        if col_name not in current_headers:
            current_headers.append(col_name)
            added.append(col_name)

    if added:
        end_col = len(current_headers)
        end_a1 = gspread.utils.rowcol_to_a1(1, end_col)
        end_col_letter = ''.join([c for c in end_a1 if c.isalpha()])
        sheet.update(f"A1:{end_col_letter}1", [current_headers])

    return {
        "headers_total": len(current_headers),
        "added": added,
        "headers": current_headers,
    }


def ensure_tender_manager_columns():
    sheet = get_sheet()
    headers = sheet.row_values(1)

    if not headers:
        ensure_sheet_columns()

    current_headers = sheet.row_values(1)
    added = []

    for col_name in TENDER_MANAGER_COLUMNS:
        if col_name not in current_headers:
            current_headers.append(col_name)
            added.append(col_name)

    if added:
        end_col = len(current_headers)
        end_a1 = gspread.utils.rowcol_to_a1(1, end_col)
        end_col_letter = ''.join([c for c in end_a1 if c.isalpha()])
        sheet.update(f"A1:{end_col_letter}1", [current_headers])

    return {
        "headers_total": len(current_headers),
        "added": added,
        "headers": current_headers,
    }


def get_header_index_map(headers):
    result = {}
    for idx, h in enumerate(headers, start=1):
        if h:
            result[h.strip()] = idx
    return result


def update_row_analytics(sheet, row_number, headers_map, analytics):
    cells = []

    for col_name, value in analytics.items():
        col_idx = headers_map.get(col_name)
        if col_idx:
            cells.append(gspread.Cell(row_number, col_idx, value))

    if cells:
        sheet.update_cells(cells, value_input_option="USER_ENTERED")

    return len(cells)


def analyze_uzex_for_sheet(site, title, url):
    empty = {
        "Логистика": "",
        "Приоритет": "",
        "Приоритет ": "",
        "Риск": "",
        "Риск ": "",
        "Шанс победы": "",
        "Заказчик": "",
        "Сумма": "",
        "Валюта": "",
        "Оплата": "",
        "Срок оплаты": "",
        "Срок оказания услуг": "",
        "Маршрут": "",
        "Тип транспорта": "",
        "Требования": "",
        "Документы нужны": "",
        "Предупреждения по документам": "",
        "Ответственные задачи": "",
        "Рекомендация AI": "",
    }

    if site != "UZEX":
        empty["Документы нужны"] = detect_required_documents_from_text(title, "", None)["summary"]
        return empty

    lot_id = extract_lot_id_from_url(url)
    if not lot_id:
        empty["Рекомендация AI"] = "Не удалось определить ID лота"
        return empty

    try:
        trade = get_uzex_lot_data(lot_id)
        budget_products = safe_json_loads(trade.get("budget_products"), [])

        delivery_term = ""
        if budget_products:
            delivery_term = budget_products[0].get("Delivery_Term", "") or ""

        quality = quality_decision_from_trade(trade, title)
        smart = smart_api_analysis_from_trade(trade, title)

        doc_list = smart.get("required_documents", [])
        docs_text = "; ".join([d.get("document", "") for d in doc_list[:20]])
        warnings_text = "; ".join(smart.get("document_warnings", [])[:10])
        responsible_text = "; ".join([f"{x.get('role')}: {x.get('task')}" for x in smart.get("responsible_tasks", [])])

        return {
            "Логистика": smart.get("logistics", quality.get("logistics", "")),
            "Приоритет": smart.get("priority", quality.get("priority", "")),
            "Приоритет ": smart.get("priority", quality.get("priority", "")),
            "Риск": smart.get("risk", quality.get("risk", "")),
            "Риск ": smart.get("risk", quality.get("risk", "")),
            "Шанс победы": smart.get("win_chance", quality.get("win_chance", "")),
            "Заказчик": trade.get("customer_name", "") or "",
            "Сумма": trade.get("start_cost", "") or "",
            "Валюта": trade.get("currency_codeabc", "") or trade.get("currency_name", "") or "",
            "Оплата": trade.get("payment_type_name", "") or "",
            "Срок оплаты": trade.get("term_payment_days", "") or "",
            "Срок оказания услуг": delivery_term,
            "Маршрут": smart.get("route_text", ""),
            "Тип транспорта": smart.get("transport_text", ""),
            "Требования": short_requirements_from_trade(trade),
            "Документы нужны": docs_text,
            "Предупреждения по документам": warnings_text,
            "Ответственные задачи": responsible_text,
            "Рекомендация AI": smart.get("decision", ""),
        }

    except Exception as e:
        empty["Рекомендация AI"] = "Ошибка анализа UZEX API: " + str(e)[:180]
        return empty


def make_row_by_headers(headers, base_values, analytics):
    base_map = {
        "Дата": base_values.get("Дата", ""),
        "Отправил": base_values.get("Отправил", ""),
        "Ссылка": base_values.get("Ссылка", ""),
        "Источник": base_values.get("Источник", ""),
        "Статус": base_values.get("Статус", ""),
        "Приоритет": base_values.get("Приоритет", ""),
        "Приоритет ": base_values.get("Приоритет", ""),
        "AI анализ": base_values.get("AI анализ", ""),
        "Комментарий": base_values.get("Комментарий", ""),
    }

    row = []
    for h in headers:
        clean_h = h.strip()
        if h in base_map:
            row.append(base_map[h])
        elif clean_h in analytics:
            row.append(analytics[clean_h])
        else:
            row.append("")

    return row


def tender_exists(url):
    try:
        sheet = get_sheet()
        urls_col_c = sheet.col_values(3)
        urls_col_d = sheet.col_values(4)
        return url in urls_col_c or url in urls_col_d
    except Exception as e:
        print("CHECK ERROR:", e)
        return False


def save_to_sheet(site, title, url):
    try:
        if tender_exists(url):
            return False

        ensure_sheet_columns()
        ensure_tender_manager_columns()
        sheet = get_sheet()
        headers = sheet.row_values(1)

        analytics = analyze_uzex_for_sheet(site, title, url)

        base_values = {
            "Дата": datetime.now().strftime("%d.%m.%Y %H:%M"),
            "Отправил": "AI Agent",
            "Ссылка": url,
            "Источник": site,
            "Статус": "Новый",
            "Приоритет": analytics.get("Приоритет", "Средний"),
            "AI анализ": "Cargo V22: AI анализ + чек-лист документов",
            "Комментарий": title,
        }

        row = make_row_by_headers(headers, base_values, analytics)
        sheet.append_row(row)

        return True

    except Exception as e:
        print("GOOGLE SHEETS ERROR:", e)
        return False


def format_document_checklist_for_telegram(documents, warnings, responsible):
    text = ""

    if documents:
        text += "\n📋 Документы для подготовки:\n"
        for i, doc in enumerate(documents[:12], start=1):
            text += f"{i}. {doc.get('document')}\n"

    if warnings:
        text += "\n⚠️ Важные предупреждения:\n"
        for w in warnings[:6]:
            text += f"- {w}\n"

    if responsible:
        text += "\n👥 Кому что сделать:\n"
        for r in responsible[:4]:
            text += f"- {r.get('role')}: {r.get('task')}\n"

    return text


def format_tender_message(tender):
    site = tender.get("site", "")
    title = tender.get("title", "")
    url = tender.get("url", "")

    analytics = analyze_uzex_for_sheet(site, title, url)
    smart = smart_api_analysis_for_url(site, title, url)

    customer = analytics.get("Заказчик") or "-"
    amount = analytics.get("Сумма") or "-"
    currency = analytics.get("Валюта") or "-"
    payment = analytics.get("Оплата") or "-"
    payment_days = analytics.get("Срок оплаты") or "-"
    delivery_term = analytics.get("Срок оказания услуг") or "-"
    route_text = analytics.get("Маршрут") or smart.get("route_text") or "-"
    transport_text = analytics.get("Тип транспорта") or smart.get("transport_text") or "-"
    risk = smart.get("risk") or analytics.get("Риск") or analytics.get("Риск ") or "-"
    win_chance = smart.get("win_chance") or analytics.get("Шанс победы") or "-"
    priority = smart.get("priority") or analytics.get("Приоритет") or "-"
    recommendation = smart.get("decision") or analytics.get("Рекомендация AI") or "Проверить вручную"
    requirements = smart.get("requirements_short") or analytics.get("Требования") or ""
    score = smart.get("score", "-")
    reason_text = smart.get("reason_text") or "-"
    risk_text = smart.get("risk_text") or "-"
    document_note = smart.get("document_note") or ""

    documents = smart.get("required_documents", [])
    warnings = smart.get("document_warnings", [])
    responsible = smart.get("responsible_tasks", [])

    if delivery_term not in ["-", ""]:
        delivery_term = str(delivery_term)
        if delivery_term.isdigit():
            delivery_term = delivery_term + " дней"

    message = (
        f"🚚 Новый логистический тендер\n\n"
        f"📌 Источник: {site}\n"
        f"🧠 AI Score: {score}/100\n"
        f"🎯 Приоритет: {priority}\n\n"
        f"📋 {title}\n\n"
        f"🏢 Заказчик: {customer}\n"
        f"💰 Сумма: {amount}\n"
        f"💵 Валюта: {currency}\n"
        f"💳 Оплата: {payment}\n"
        f"⏳ Срок оплаты: {payment_days} дней\n"
        f"📅 Срок оказания: {delivery_term}\n\n"
        f"🗺 Маршрут: {route_text}\n"
        f"🚛 Транспорт: {transport_text}\n\n"
        f"📈 Шанс победы: {win_chance}\n"
        f"⚠️ Риск: {risk}\n\n"
        f"✅ Почему интересно:\n{reason_text}\n\n"
        f"⚠️ Что проверить:\n{risk_text}\n\n"
        f"🤖 AI рекомендация:\n{recommendation}\n"
    )

    if requirements:
        message += f"\n📄 Требования/API кратко:\n{requirements[:500]}\n"

    message += format_document_checklist_for_telegram(documents, warnings, responsible)

    if document_note:
        message += f"\n📎 Документы: {document_note}\n"

    message += f"\n🔗 {url}"

    return message[:3900]


def build_file_url_candidates(file_path):
    if not file_path:
        return []

    file_path = str(file_path).strip()
    candidates = []

    if file_path.startswith("http://") or file_path.startswith("https://"):
        candidates.append(file_path)
    else:
        clean = file_path.lstrip("/")

        host_variants = [
            "https://etender.uzex.uz/",
            "https://apietender.uzex.uz/",
        ]

        path_variants = [clean]

        if clean.startswith("files/"):
            path_variants.append("tender/user-files/" + clean[len("files/"):])
            path_variants.append("user-files/" + clean[len("files/"):])

        if clean.startswith("tender/user-files/"):
            tail = clean[len("tender/user-files/"):]
            path_variants.append("files/" + tail)
            path_variants.append("user-files/" + tail)

        if clean.startswith("user-files/"):
            tail = clean[len("user-files/"):]
            path_variants.append("files/" + tail)
            path_variants.append("tender/user-files/" + tail)

        for host in host_variants:
            for path in path_variants:
                candidates.append(host + path)

    result = []
    seen = set()
    for u in candidates:
        if u not in seen:
            seen.add(u)
            result.append(u)

    return result


def is_html_response(content, content_type=""):
    head = (content or b"")[:500].lower()
    ctype = (content_type or "").lower()
    return (
        "text/html" in ctype
        or b"<!doctype html" in head
        or b"<html" in head
        or b"<app-root" in head
    )


def detect_file_kind(content, content_type="", url=""):
    content = content or b""
    ctype = (content_type or "").lower()
    url_low = (url or "").lower()

    if is_html_response(content, ctype):
        return "html"
    if content.startswith(b"%PDF") or "application/pdf" in ctype or url_low.endswith(".pdf"):
        return "pdf" if content.startswith(b"%PDF") else "maybe_pdf"
    if content.startswith(b"PK"):
        if url_low.endswith(".docx") or "wordprocessingml" in ctype:
            return "docx"
        if url_low.endswith(".xlsx") or "spreadsheetml" in ctype:
            return "xlsx"
        return "zip"
    return "unknown"


def download_file_with_fallback(file_path):
    attempts = []

    for file_url in build_file_url_candidates(file_path):
        try:
            r = requests.get(
                file_url,
                headers={
                    **get_headers(),
                    "Accept": "application/pdf,application/vnd.openxmlformats-officedocument.wordprocessingml.document,application/vnd.openxmlformats-officedocument.spreadsheetml.sheet,application/octet-stream,*/*",
                    "Referer": "https://etender.uzex.uz/",
                },
                timeout=30,
                allow_redirects=True,
            )
            content_type = r.headers.get("content-type", "")
            kind = detect_file_kind(r.content, content_type, file_url)

            info = {
                "url": file_url,
                "status_code": r.status_code,
                "content_type": content_type,
                "size": len(r.content or b""),
                "detected_kind": kind,
            }

            if kind == "html":
                info["rejected"] = "html_page_not_document"
                info["html_start"] = (r.text or "")[:300]

            attempts.append(info)

            if r.status_code == 200 and r.content and kind in ["pdf", "docx", "xlsx", "zip"]:
                return r.content, file_url, attempts

            if r.status_code == 200 and kind in ["maybe_pdf", "unknown"]:
                attempts[-1]["rejected"] = "not_real_document_bytes"

        except Exception as e:
            attempts.append({
                "url": file_url,
                "error": str(e),
            })

    raise Exception("No real document downloaded. Attempts: " + json.dumps(attempts, ensure_ascii=False)[:1800])


def read_pdf_from_bytes(file_bytes, max_pages=12):
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text_parts = []

    for i, page in enumerate(reader.pages[:max_pages]):
        try:
            page_text = page.extract_text() or ""
            if page_text:
                text_parts.append(page_text)
        except Exception as e:
            text_parts.append(f"[PDF PAGE {i + 1} READ ERROR: {e}]")

    return clean_text("\n".join(text_parts))


def read_docx_from_bytes(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for p in doc.paragraphs:
        txt = clean_text(p.text)
        if txt:
            text_parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(clean_text(cell.text) for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return clean_text("\n".join(text_parts))


def pick_text_snippet(text, keywords, limit=700):
    t = text or ""
    low = t.lower()

    for kw in keywords:
        pos = low.find(kw.lower())
        if pos >= 0:
            start = max(0, pos - 250)
            end = min(len(t), pos + limit)
            return clean_text(t[start:end])

    return clean_text(t[:limit])


@app.get("/")
def home():
    return {"status": "AI Tender Agent Cargo V22 Document Checklist is running"}


@app.head("/")
def head_home():
    return {}


@app.get("/version")
def version():
    return {
        "version": "cargo_v22_document_checklist",
        "status": "running"
    }


@app.get("/health")
def health():
    result = {"status": "ok", "telegram": False, "google_sheets": False, "errors": []}

    try:
        tg = requests.get(f"https://api.telegram.org/bot{BOT_TOKEN}/getMe", timeout=8)
        result["telegram"] = tg.status_code == 200
    except Exception as e:
        result["errors"].append("Telegram error: " + str(e))

    try:
        sheet = get_sheet()
        sheet.row_values(1)
        result["google_sheets"] = True
    except Exception as e:
        result["errors"].append("Google Sheets error: " + str(e))

    if result["errors"]:
        result["status"] = "warning"

    return result


@app.get("/test_filter")
def test_filter():
    return {
        "Услуга по перевозке грузов": is_real_cargo_tender("Услуга по перевозке грузов", "https://etender.uzex.uz/lot/123"),
        "Оказание транспортных услуг": is_real_cargo_tender("Оказание транспортных услуг", "https://xarid.uzex.uz/tender/456"),
        "Транспортно-экспедиционные услуги": is_real_cargo_tender("Транспортно-экспедиционные услуги", "https://etender.uzex.uz/lot/999"),
        "Лабораторное оборудование": is_real_cargo_tender("Лабораторное оборудование", "https://www.tenderweek.com/tender-35921"),
        "Закупка арматуры для АЭС": is_real_cargo_tender("Закупка арматуры для АЭС", "https://www.tenderweek.com/tender-35911"),
        "Доставка товара автотранспортом": is_real_cargo_tender("Доставка товара автотранспортом", "https://www.tenderweek.com/tender-99999"),
        "xalqaro yuk tashish xizmati": is_real_cargo_tender("xalqaro yuk tashish xizmati", "https://xt-xarid.uz/tender/999"),
        "Автоматизация кредитного конвейера": is_real_cargo_tender("Предоставление услуг по автоматизации кредитного конвейера", "https://xt-xarid.uz/tender/111"),
    }


@app.get("/analyze_doc_test")
def analyze_doc_test():
    return {
        "status": "ok",
        "version": "document_analyzer_v22",
        "pdf_reader": True,
        "docx_reader": True,
        "xlsx_reader": True,
        "document_checklist": True,
        "modules": {
            "PyPDF2": True,
            "python_docx": True,
            "openpyxl": True
        }
    }


@app.get("/setup_sheet_columns")
def setup_sheet_columns():
    try:
        result = ensure_sheet_columns()
        manager = ensure_tender_manager_columns()
        return {
            "status": "ok",
            "version": "sheet_setup_v22",
            "message": "Google Sheets columns checked and updated",
            "analytics": result,
            "manager": manager,
        }
    except Exception as e:
        return {
            "status": "error",
            "version": "sheet_setup_v22",
            "error": str(e),
        }


@app.get("/setup_tender_manager_columns")
def setup_tender_manager_columns():
    try:
        ensure_sheet_columns()
        result = ensure_tender_manager_columns()

        return {
            "status": "ok",
            "version": "tender_manager_v22",
            "message": "Tender manager columns checked and updated",
            **result,
        }

    except Exception as e:
        return {
            "status": "error",
            "version": "tender_manager_v22",
            "error": str(e),
        }


@app.get("/tender_manager_status")
def tender_manager_status():
    try:
        sheet = get_sheet()
        headers = sheet.row_values(1)

        missing = []
        for col in EXTRA_SHEET_COLUMNS + TENDER_MANAGER_COLUMNS:
            if col not in headers:
                missing.append(col)

        return {
            "status": "ok" if not missing else "warning",
            "version": "tender_manager_v22",
            "headers_total": len(headers),
            "missing_columns": missing,
            "manager_columns": TENDER_MANAGER_COLUMNS,
            "analytics_columns": EXTRA_SHEET_COLUMNS,
        }

    except Exception as e:
        return {
            "status": "error",
            "version": "tender_manager_v22",
            "error": str(e),
        }


@app.get("/debug_sources")
def debug_sources():
    result = {
        "version": "cargo_v22_document_checklist",
        "Tenderweek": 0,
        "UZEX": 0,
        "XT-Xarid": 0,
        "total": 0,
        "errors": [],
    }

    try:
        tw = parse_tenderweek()
        result["Tenderweek"] = len(tw)
    except Exception as e:
        result["errors"].append("Tenderweek error: " + str(e))

    try:
        uzex = parse_uzex()
        result["UZEX"] = len(uzex)
    except Exception as e:
        result["errors"].append("UZEX error: " + str(e))

    try:
        xt = parse_xt_xarid()
        result["XT-Xarid"] = len(xt)
    except Exception as e:
        result["errors"].append("XT-Xarid error: " + str(e))

    result["total"] = result["Tenderweek"] + result["UZEX"] + result["XT-Xarid"]
    return result


@app.get("/debug_items")
def debug_items():
    all_items = []

    for source_name, parser in [
        ("Tenderweek", parse_tenderweek),
        ("UZEX", parse_uzex),
        ("XT-Xarid", parse_xt_xarid),
    ]:
        try:
            result = parser()
            for item in result[:10]:
                all_items.append({
                    "site": source_name,
                    "title": item.get("title"),
                    "url": item.get("url"),
                    "reason": item.get("reason"),
                })
        except Exception as e:
            all_items.append({
                "site": source_name,
                "error": str(e),
            })

    return {
        "version": "cargo_v22_document_checklist",
        "count": len(all_items),
        "items": all_items[:30],
    }


@app.get("/debug_raw_candidates")
def debug_raw_candidates():
    all_items = []

    for source_name, parser in [
        ("Tenderweek", parse_tenderweek),
        ("UZEX", parse_uzex),
        ("XT-Xarid", parse_xt_xarid),
    ]:
        try:
            result = parser(raw=True)
            for item in result[:25]:
                all_items.append(item)
        except Exception as e:
            all_items.append({
                "site": source_name,
                "error": str(e),
            })

    accepted = [x for x in all_items if x.get("accepted") is True]
    rejected = [x for x in all_items if x.get("accepted") is False]

    return {
        "version": "cargo_v22_document_checklist",
        "total_candidates_sample": len(all_items),
        "accepted_sample": len(accepted),
        "rejected_sample": len(rejected),
        "items": all_items[:75],
    }


@app.get("/analyze_uzex_lot")
def analyze_uzex_lot(lot_id: str):
    try:
        trade = get_uzex_trade(lot_id)

        budget_products = extract_budget_products(trade)
        product_name = ""
        product_description = ""
        category_name = ""
        delivery_term_days = None

        if budget_products:
            product_name = budget_products[0].get("Product_Name", "") or ""
            product_description = budget_products[0].get("Description", "") or ""
            category_name = budget_products[0].get("Category_Name", "") or ""
            delivery_term_days = budget_products[0].get("Delivery_Term")

        smart = smart_api_analysis_from_trade(trade, " | ".join([product_name, product_description]))

        documents = []
        file_fields = [
            ("tech_file", trade.get("tech_file_name"), trade.get("tech_file_path"), trade.get("tech_file_ext")),
            ("tech_doc_file", trade.get("tech_doc_file_name"), trade.get("tech_doc_file_path"), trade.get("tech_doc_file_ext")),
            ("contract_proform_file", trade.get("contract_proform_file_name"), trade.get("contract_proform_file_path"), trade.get("contract_proform_file_ext")),
            ("prolong_file", trade.get("prolong_file_name"), trade.get("prolong_file_path"), trade.get("prolong_file_ext")),
        ]

        combined_text = ""
        for doc_type, name, path, ext in file_fields:
            if not path:
                continue

            doc_info = {
                "type": doc_type,
                "name": name,
                "ext": ext,
                "path": path,
                "url_candidates": build_file_url_candidates(path),
                "working_url": None,
                "download_attempts": [],
                "read_status": "not_read",
                "text_preview": "",
                "note": "V22: API-анализ и чек-лист работают даже если документ закрыт прямым скачиванием",
            }

            try:
                ext_norm = (ext or name or path or "").lower()
                file_bytes, working_url, attempts = download_file_with_fallback(path)
                doc_info["working_url"] = working_url
                doc_info["download_attempts"] = attempts

                if "pdf" in ext_norm:
                    text = read_pdf_from_bytes(file_bytes)
                    doc_info["read_status"] = "ok"
                    doc_info["text_preview"] = text[:1200]
                    combined_text += "\n" + text
                elif "docx" in ext_norm:
                    text = read_docx_from_bytes(file_bytes)
                    doc_info["read_status"] = "ok"
                    doc_info["text_preview"] = text[:1200]
                    combined_text += "\n" + text
                else:
                    doc_info["read_status"] = "unsupported_ext"
            except Exception as e:
                doc_info["read_status"] = "closed_or_unavailable"
                doc_info["error"] = str(e)[:900]

            documents.append(doc_info)

        document_ok_count = sum(1 for d in documents if d.get("read_status") == "ok")

        full_text_for_docs = " ".join([
            product_name,
            product_description,
            category_name,
            smart.get("api_text_used", ""),
            combined_text
        ])

        doc_checklist = detect_required_documents_from_text(
            " | ".join([product_name, product_description]),
            full_text_for_docs,
            trade
        )

        route_transport = extract_route_and_transport(
            " | ".join([product_name, product_description]),
            full_text_for_docs
        )

        return {
            "status": "ok",
            "version": "document_analyzer_v22",
            "lot_id": lot_id,
            "api_url": f"https://apietender.uzex.uz/api/common/GetTrade/{lot_id}/0",
            "lot": {
                "id": trade.get("id"),
                "display_no": trade.get("display_no"),
                "customer_name": trade.get("customer_name"),
                "customer_tin": trade.get("customer_tin"),
                "start_date": trade.get("start_date"),
                "end_date": trade.get("end_date"),
                "start_cost": trade.get("start_cost"),
                "currency": trade.get("currency_codeabc") or trade.get("currency_name"),
                "valuation_name": trade.get("valuation_name"),
                "payment_type_name": trade.get("payment_type_name"),
                "term_payment_days": trade.get("term_payment_days"),
                "delivery_term_days": delivery_term_days,
                "product_name": product_name,
                "product_description": product_description,
                "category_name": category_name,
            },
            "smart_api_analysis": {
                "ai_score": smart["score"],
                "priority": smart["priority"],
                "win_chance": smart["win_chance"],
                "risk": smart["risk"],
                "logistics": smart["logistics"],
                "decision": smart["decision"],
                "reasons": smart["reasons"],
                "risks": smart["risks"],
                "reason_text": smart["reason_text"],
                "risk_text": smart["risk_text"],
                "requirements_short": smart["requirements_short"],
                "document_note": smart["document_note"],
                "route": route_transport,
            },
            "document_checklist": doc_checklist,
            "documents_count": len(documents),
            "documents_read_ok": document_ok_count,
            "documents": documents,
            "analysis": {
                "priority": smart["priority"],
                "win_chance": smart["win_chance"],
                "risk": smart["risk"],
                "logistics": smart["logistics"],
                "decision": smart["decision"],
                "route": route_transport["route_text"],
                "transport": route_transport["transport_text"],
                "requirements_snippet": pick_text_snippet(
                    combined_text,
                    ["требования к участнику", "тягач", "goldhofer", "gantry", "требования к персоналу"]
                ) or smart["requirements_short"],
                "payment_snippet": pick_text_snippet(
                    combined_text,
                    ["условия оплаты", "оплата", "payment", "30 календарных дней", "15 дней"]
                ) or clean_text(f"{trade.get('payment_type_name') or ''}; срок оплаты: {trade.get('term_payment_days') or ''} дней"),
                "summary": (
                    f"AI Score: {smart['score']}/100. "
                    f"Заказчик: {trade.get('customer_name')}. "
                    f"Предмет: {product_name}. "
                    f"Стартовая цена: {trade.get('start_cost')} {trade.get('currency_codeabc') or trade.get('currency_name')}. "
                    f"Документов к подготовке: {doc_checklist['documents_count']}. "
                    f"Рекомендация: {smart['decision']}"
                ),
            },
        }

    except Exception as e:
        return {
            "status": "error",
            "version": "document_analyzer_v22",
            "lot_id": lot_id,
            "error": str(e),
        }


@app.get("/document_checklist")
def document_checklist(lot_id: str):
    try:
        trade = get_uzex_trade(lot_id)
        budget_products = extract_budget_products(trade)

        title = ""
        text_parts = []

        if budget_products:
            first = budget_products[0]
            title = clean_text(first.get("Product_Name", ""))
            text_parts.append(clean_text(first.get("Description", "")))
            text_parts.append(clean_text(first.get("Category_Name", "")))

        text_parts.append(clean_text(trade.get("technical_description", "")))
        text_parts.append(clean_text(trade.get("customer_name", "")))

        checklist = detect_required_documents_from_text(title, " ".join(text_parts), trade)

        return {
            "status": "ok",
            "version": "document_checklist_v22",
            "lot_id": lot_id,
            "title": title,
            **checklist,
        }
    except Exception as e:
        return {
            "status": "error",
            "version": "document_checklist_v22",
            "lot_id": lot_id,
            "error": str(e),
        }


@app.get("/backfill_existing_tenders")
def backfill_existing_tenders(limit: int = 50):
    try:
        ensure_sheet_columns()
        ensure_tender_manager_columns()

        sheet = get_sheet()
        all_values = sheet.get_all_values()

        if not all_values:
            return {
                "status": "warning",
                "version": "backfill_v22",
                "message": "Sheet is empty",
            }

        headers = all_values[0]
        headers_map = get_header_index_map(headers)

        if "Ссылка" not in headers_map:
            return {
                "status": "error",
                "version": "backfill_v22",
                "error": "Required column not found: Ссылка",
                "headers": headers,
            }

        processed = 0
        updated = 0
        skipped = 0
        errors = []

        link_col = headers_map["Ссылка"]

        for row_idx, row in enumerate(all_values[1:], start=2):
            if processed >= limit:
                break

            link = row[link_col - 1] if len(row) >= link_col else ""

            if not link or "etender.uzex.uz/lot/" not in link:
                skipped += 1
                continue

            processed += 1

            try:
                analytics = analyze_uzex_for_sheet("UZEX", "Backfill UZEX lot", link)
                update_row_analytics(sheet, row_idx, headers_map, analytics)
                updated += 1

            except Exception as e:
                errors.append({
                    "row": row_idx,
                    "link": link,
                    "error": str(e)[:250],
                })

        return {
            "status": "ok",
            "version": "backfill_v22",
            "processed": processed,
            "updated": updated,
            "skipped": skipped,
            "errors_count": len(errors),
            "errors": errors[:10],
            "message": "Existing UZEX tenders backfill with document checklist completed",
        }

    except Exception as e:
        return {
            "status": "error",
            "version": "backfill_v22",
            "error": str(e),
        }


@app.get("/quality_backfill_existing_tenders")
def quality_backfill_existing_tenders(limit: int = 100):
    return backfill_existing_tenders(limit=limit)


@app.get("/scan")
def scan():
    print("SCAN STARTED V22")

    found_total = 0
    new_total = 0
    duplicate_total = 0
    all_tenders = []
    seen_urls = set()

    message = "📊 AI Tender Agent Cargo V22 Document Checklist Scan завершён\n\n"

    sources = [
        ("Tenderweek", parse_tenderweek),
        ("UZEX", parse_uzex),
        ("XT-Xarid", parse_xt_xarid),
    ]

    source_counts = {}

    for source_name, parser in sources:
        print("PARSING:", source_name)

        try:
            result = parser()
            source_counts[source_name] = len(result)
            all_tenders.extend(result)
            message += f"{source_name}: найдено услуг по перевозке {len(result)}\n"
        except Exception as e:
            source_counts[source_name] = 0
            message += f"{source_name}: ERROR\n"
            print(f"{source_name} ERROR:", e)

    message += "\n"

    for tender in all_tenders:
        url = tender.get("url")
        title = tender.get("title")

        if not url or url in seen_urls:
            continue

        seen_urls.add(url)

        if not is_real_cargo_tender(title, url):
            continue

        found_total += 1

        saved = save_to_sheet(tender["site"], title, url)

        if saved:
            new_total += 1
            send_telegram(format_tender_message(tender))
        else:
            duplicate_total += 1

    message += (
        f"Всего найдено: {found_total}\n"
        f"Новых сохранено: {new_total}\n"
        f"Дубликатов пропущено: {duplicate_total}\n\n"
        f"V22: добавлен чек-лист документов и задачи для сотрудников."
    )

    send_telegram(message)

    return {
        "status": "success",
        "version": "cargo_v22_document_checklist",
        "sources": source_counts,
        "found_total": found_total,
        "new_total": new_total,
        "duplicates": duplicate_total,
    }


@app.head("/scan")
def scan_head():
    return {}


@app.post("/webhook")
def webhook():
    return scan()


@app.get("/webhook")
def webhook_get():
    return scan()


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=int(os.getenv("PORT", 10000)),
